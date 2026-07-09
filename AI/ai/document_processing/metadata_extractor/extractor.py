import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
import fitz
import docx
import openpyxl
from loguru import logger


class MetadataExtractor:
    """Metadata extractor for engineering and technical document files."""

    def extract(self, file_path: Path) -> Dict[str, Any]:
        """Extracts system and content metadata from a document file.

        Args:
            file_path: Path to the target document.

        Returns:
            Dict[str, Any]: Dict containing parsed metadata matching DocumentMetadata schema.
        """
        logger.info("Extracting metadata for file: {}", file_path.name)

        # Baseline file system metadata
        stats = file_path.stat()
        metadata = {
            "file_name": file_path.name,
            "file_extension": file_path.suffix.lower(),
            "file_size_bytes": stats.st_size,
            "total_pages": 1,
            "author": None,
            "creation_timestamp": None,
            "processing_timestamp": datetime.utcnow().isoformat() + "Z",
            "sheet_names": None,
            "custom_metadata": {}
        }

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            self._extract_pdf_metadata(file_path, metadata)
        elif ext == ".docx":
            self._extract_docx_metadata(file_path, metadata)
        elif ext in [".xlsx", ".xls"]:
            self._extract_excel_metadata(file_path, metadata)

        return metadata

    def _extract_pdf_metadata(self, file_path: Path, metadata: Dict[str, Any]) -> None:
        """Extracts metadata from PDF document container."""
        try:
            doc = fitz.open(str(file_path))
            metadata["total_pages"] = len(doc)
            pdf_meta = doc.metadata
            if pdf_meta:
                metadata["author"] = pdf_meta.get("author") or None
                metadata["creation_timestamp"] = pdf_meta.get("creationDate") or None
                metadata["custom_metadata"] = {
                    "creator": pdf_meta.get("creator") or "",
                    "producer": pdf_meta.get("producer") or "",
                    "subject": pdf_meta.get("subject") or "",
                    "title": pdf_meta.get("title") or ""
                }
            doc.close()
        except Exception as e:
            logger.warning("Could not extract detailed PDF metadata for {}: {}", file_path.name, str(e))

    def _extract_docx_metadata(self, file_path: Path, metadata: Dict[str, Any]) -> None:
        """Extracts metadata from DOCX document container."""
        try:
            doc = docx.Document(str(file_path))
            props = doc.core_properties
            if props:
                metadata["author"] = props.author or None
                if props.created:
                    metadata["creation_timestamp"] = props.created.isoformat() + "Z"
            # Extract additional technical metadata
            metadata["custom_metadata"] = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        except Exception as e:
            logger.warning("Could not extract DOCX metadata for {}: {}", file_path.name, str(e))

    def _extract_excel_metadata(self, file_path: Path, metadata: Dict[str, Any]) -> None:
        """Extracts metadata from Excel workbook container."""
        try:
            wb = openpyxl.load_workbook(filename=str(file_path), read_only=True)
            metadata["sheet_names"] = wb.sheetnames
            metadata["total_pages"] = len(wb.sheetnames)
            props = wb.properties
            if props:
                metadata["author"] = props.creator or None
                if props.created:
                    metadata["creation_timestamp"] = props.created.isoformat() + "Z"
            wb.close()
        except Exception as e:
            logger.warning("Could not extract Excel metadata for {}: {}", file_path.name, str(e))
