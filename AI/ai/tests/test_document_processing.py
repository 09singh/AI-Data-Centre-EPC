import pytest
from pathlib import Path
import fitz
import docx
import openpyxl

from ai.document_processing.exceptions import UnsupportedDocumentError, EmptyDocumentError
from ai.document_processing.pdf_reader.reader import PDFReader
from ai.document_processing.docx_reader.reader import DOCXReader
from ai.document_processing.excel_reader.reader import ExcelReader
from ai.document_processing.text_cleaner.cleaner import TextCleaner
from ai.document_processing.metadata_extractor.extractor import MetadataExtractor
from ai.document_processing.chunking.chunker import Chunker
from ai.document_processing.pipeline import DocumentProcessingPipeline


@pytest.fixture
def temp_pdf(tmp_path) -> Path:
    """Fixture to generate a mock PDF document dynamically using PyMuPDF."""
    path = tmp_path / "test.pdf"
    doc = fitz.open()
    # Create Page 1
    page1 = doc.new_page()
    page1.insert_text((50, 50), "1. INTRODUCTION\nThis is the first page of the EPC engineering specifications.")
    # Create Page 2
    page2 = doc.new_page()
    page2.insert_text((50, 50), "2. ELECTRICAL REQS\nThis page outlines data centre design calculations. Power is 100 MW.")
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def temp_docx(tmp_path) -> Path:
    """Fixture to generate a mock DOCX document dynamically using python-docx."""
    path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("1. SYSTEM DESCRIPTION")
    doc.add_paragraph("This document outlines the mechanical piping layout of the data centre cooling loop. Voltage is 10 kV.")

    # Interleave a table structure
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Equipment Tag"
    hdr_cells[1].text = "Design Capacity"
    row_cells = table.rows[1].cells
    row_cells[0].text = "CRAH-01"
    row_cells[1].text = "150 kW"

    doc.save(str(path))
    return path


@pytest.fixture
def temp_xlsx(tmp_path) -> Path:
    """Fixture to generate a mock Excel workbook dynamically using openpyxl."""
    path = tmp_path / "test.xlsx"
    wb = openpyxl.Workbook()

    # Active sheet (Sheet 1)
    ws1 = wb.active
    ws1.title = "LoadsSummary"
    ws1.append(["Load ID", "Facility Zone", "Capacity MW"])
    ws1.append(["L-101", "Data Hall A", "12.5"])
    ws1.append(["L-102", "Data Hall B", "12.5"])

    # Sheet 2
    ws2 = wb.create_sheet(title="Generators")
    ws2.append(["Generator ID", "Rating kV", "Fuel Configuration"])
    ws2.append(["GEN-001", "11", "Diesel Dual-Fuel"])

    wb.save(str(path))
    return path


def test_pdf_reader(temp_pdf):
    """Tests page-wise extraction, page count, and contents of the PDF reader."""
    reader = PDFReader()
    pages = reader.extract_text_by_page(temp_pdf)
    assert len(pages) == 2
    assert pages[0]["page_number"] == 1
    assert "INTRODUCTION" in pages[0]["text"]
    assert pages[1]["page_number"] == 2
    assert "ELECTRICAL REQS" in pages[1]["text"]
    assert "100 MW" in pages[1]["text"]


def test_docx_reader(temp_docx):
    """Tests paragraphs and tables extraction from a DOCX file in order."""
    reader = DOCXReader()
    elements = reader.extract_content(temp_docx)
    assert len(elements) >= 3
    # First block is paragraph text
    assert elements[0]["type"] == "paragraph"
    assert elements[0]["text"] == "1. SYSTEM DESCRIPTION"
    # Third block is the table block converted to markdown table
    assert elements[2]["type"] == "table"
    assert "Equipment Tag" in elements[2]["text"]
    assert "CRAH-01" in elements[2]["text"]
    assert "|" in elements[2]["text"]


def test_excel_reader(temp_xlsx):
    """Tests sheet names extraction and tabular layouts formatting to Markdown tables."""
    reader = ExcelReader()
    sheets = reader.extract_sheets_content(temp_xlsx)
    assert len(sheets) == 2
    assert sheets[0]["sheet_name"] == "LoadsSummary"
    assert "LoadsSummary" in sheets[0]["text"] or "Load ID" in sheets[0]["text"]
    assert sheets[1]["sheet_name"] == "Generators"
    assert "GEN-001" in sheets[1]["text"]


def test_text_cleaner():
    """Tests text cleanup, Unicode normalizations, spacing normalization, and engineering units preservation."""
    cleaner = TextCleaner()
    raw = "Specification   1.1.2\n\n\nTotal rating:\u00a0100\u00a0MW  "
    cleaned = cleaner.clean(raw)
    assert cleaned == "Specification 1.1.2\n\nTotal rating: 100 MW"


def test_metadata_extractor(temp_pdf, temp_xlsx):
    """Tests file metadata extraction for system metadata, author properties, and page layouts."""
    extractor = MetadataExtractor()
    # Verify PDF metadata properties
    pdf_meta = extractor.extract(temp_pdf)
    assert pdf_meta["file_name"] == "test.pdf"
    assert pdf_meta["file_extension"] == ".pdf"
    assert pdf_meta["total_pages"] == 2

    # Verify Excel metadata properties
    xlsx_meta = extractor.extract(temp_xlsx)
    assert xlsx_meta["file_name"] == "test.xlsx"
    assert xlsx_meta["total_pages"] == 2
    assert xlsx_meta["sheet_names"] == ["LoadsSummary", "Generators"]


def test_chunker():
    """Tests text partition boundaries, heading preservation, page mappings, and unique chunk IDs."""
    chunker = Chunker(chunk_size=150, chunk_overlap=30)
    blocks = [
        {"text": "1. SCOPE OF WORK\nDesigning the backup battery systems.", "page_number": 1, "section_heading": "SCOPE OF WORK"},
        {"text": "2. SYSTEM TOPOLOGY\nProviding redundant UPS feeds to data halls.", "page_number": 2, "section_heading": "SYSTEM TOPOLOGY"}
    ]
    chunks = chunker.chunk_document("testhash", blocks)
    assert len(chunks) > 0
    for chunk in chunks:
        assert chunk.chunk_id is not None
        assert len(chunk.page_numbers) > 0
        assert chunk.section_heading in ["SCOPE OF WORK", "SYSTEM TOPOLOGY"]


def test_pipeline(temp_pdf, temp_docx, temp_xlsx):
    """Tests pipeline end-to-end processing execution for all three file formats."""
    pipeline = DocumentProcessingPipeline()

    # Process PDF
    pdf_res = pipeline.process(temp_pdf, chunk_size=200, chunk_overlap=50)
    assert pdf_res.metadata.file_name == "test.pdf"
    assert len(pdf_res.chunks) > 0
    assert "INTRODUCTION" in pdf_res.full_extracted_text

    # Process DOCX
    docx_res = pipeline.process(temp_docx, chunk_size=200, chunk_overlap=50)
    assert docx_res.metadata.file_name == "test.docx"
    assert len(docx_res.chunks) > 0
    assert "CRAH-01" in docx_res.full_extracted_text

    # Process Excel
    xlsx_res = pipeline.process(temp_xlsx, chunk_size=200, chunk_overlap=50)
    assert xlsx_res.metadata.file_name == "test.xlsx"
    assert xlsx_res.metadata.sheet_names == ["LoadsSummary", "Generators"]
    assert len(xlsx_res.chunks) > 0


def test_pipeline_errors(tmp_path):
    """Tests pipeline validation exceptions for invalid and unsupported input formats."""
    pipeline = DocumentProcessingPipeline()
    unsupported_file = tmp_path / "test.txt"
    unsupported_file.write_text("plain text")

    with pytest.raises(UnsupportedDocumentError):
        pipeline.process(unsupported_file)
